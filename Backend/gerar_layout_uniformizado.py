#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gerador_layout_uniformizado.py - Motor de Geracao de Documentos no Layout Canónico Uniformizado (DOCX + PDF + MD).
Aplica os padroes visuais e redatoriais fixados no MANUAL_ESTILO_E_LAYOUT_UNIFORMIZADO.md.
Zero emojis conforme PROTOCOL.md e AGENTS.md.
"""

import os
import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm

DEV_ROOT = Path(r"C:\Users\Yokozuna\Dev")
OUTPUT_DIR = DEV_ROOT / "OUTPUT_CENTRALIZADO" / "05_PDFS_GERADOS_PARA_IMPRESSAO"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_OUT = OUTPUT_DIR / "MODELO_PADRAO_PECA_JUDICIAL.docx"
PDF_OUT = OUTPUT_DIR / "MODELO_PADRAO_PECA_JUDICIAL.pdf"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def gerar_modelo_docx():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(0.8)

    # Cabecalho Duplo
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    cell_l, cell_r = header_table.rows[0].cells
    cell_l.width = Inches(4.2)
    cell_r.width = Inches(2.3)

    p_l = cell_l.paragraphs[0]
    r1 = p_l.add_run("TRIBUNAL JUDICIAL DA COMARCA DE LISBOA\n")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p_l.add_run("Juízo Central Cível de Lisboa\nPalácio da Justiça — Lisboa")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p_r.add_run("PROCESSO N.º 15547/26.0T8LSB\n")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p_r.add_run("Espécie: Ação de Reivindicação\nValor da Causa: 125.000,00 €")
    r4.font.size = Pt(8.5)
    r4.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Enderecamento
    p_j = doc.add_paragraph()
    r_j = p_j.add_run("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DO JUÍZO CENTRAL CÍVEL DE LISBOA")
    r_j.bold = True
    r_j.font.size = Pt(11)
    p_j.paragraph_format.space_after = Pt(14)

    # Titulo
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("CONTESTAÇÃO E RECONVENÇÃO\n(COM ARGUIÇÃO DE DIREITO DE RETENÇÃO E PRETERIÇÃO DE LITISCONSÓRCIO)")
    r_t.bold = True
    r_t.font.size = Pt(12.5)
    p_t.paragraph_format.space_after = Pt(16)

    # Qualificacao
    p_q = doc.add_paragraph()
    p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_q.paragraph_format.line_spacing = 1.3
    p_q.paragraph_format.space_after = Pt(10)
    p_q.add_run("NUNO MIGUEL SILVA DUARTE").bold = True
    p_q.add_run(", NIF 254048382, Réu nos autos da Ação Declarativa Comum à margem identificada em que é Autora ")
    p_q.add_run("MARIA TERESA CASTRO BANGUESES RIBEIRO").bold = True
    p_q.add_run(", vem mui respeitosamente apresentar a sua CONTESTAÇÃO, nos termos e com os fundamentos seguintes:")

    # Secao I
    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.space_before = Pt(12)
    p_s1.paragraph_format.space_after = Pt(4)
    r_s1 = p_s1.add_run("I. POR EXCEÇÃO DILATÓRIA: DA PRETERIÇÃO DE LITISCONSÓRCIO NECESSÁRIO (ART. 33.º CPC)")
    r_s1.bold = True
    r_s1.font.size = Pt(11)

    p_p1 = doc.add_paragraph()
    p_p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_p1.paragraph_format.first_line_indent = Inches(0.4)
    p_p1.paragraph_format.line_spacing = 1.3
    p_p1.add_run("1. A Autora intentou a presente ação apenas contra o Réu singular, omitindo deliberadamente a sociedade comercial ")
    p_p1.add_run("LISBON EXPERIENCE - ADMINISTRAÇÃO DE IMÓVEIS, LDA.").bold = True
    p_p1.add_run(", com a qual outorgou o Contrato de Arrendamento N.º 1195528 (referente à fração do 4.º Andar / Matriz 110661-U-231-4), conforme prova a relação oficial de contratos constante de ")
    p_p1.add_run("LISTA_CONTRATOS_TERESA.xls").bold = True
    p_p1.add_run(".")

    # Secao II
    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.space_before = Pt(12)
    p_s2.paragraph_format.space_after = Pt(4)
    r_s2 = p_s2.add_run("II. DA PROVA DOCUMENTAL E DAS CONFISSÕES ESCRITAS SOBRE CRÉDITOS E CORTE DE ÁGUA")
    r_s2.bold = True
    r_s2.font.size = Pt(11)

    p_p2 = doc.add_paragraph()
    p_p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_p2.paragraph_format.first_line_indent = Inches(0.4)
    p_p2.paragraph_format.line_spacing = 1.3
    p_p2.add_run("2. O desconhecimento do Réu quanto ao desvio de fundos e a confissão de que era titular de rendimentos a receber resultam de forma cristalina das comunicações com a gerência da Lisbon Experience:")

    # Caixas de Citacao
    def add_quote(speaker, dt, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tbl.rows[0].cells[0]
        c.width = Inches(6.0)
        set_cell_background(c, "F1F5F9")
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        r_h = p.add_run(f"[{dt}] {speaker}:\n")
        r_h.bold = True
        r_h.font.size = Pt(9.5)
        r_txt = p.add_run(f"\"{text}\"")
        r_txt.italic = True
        r_txt.font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_quote("Nuno Duarte", "2023-01-24, 17:14", "Apesar de Eu e o Filipe e o cleber termos assinado o doc, nao impede de ele lançar uma acao contra mim.")
    add_quote("Filipe Delgado", "2023-01-24, 17:15", "O único responsável sou eu. Documento foi assinado por mim como confissão de dívida e autenticado pelo advogado. Eu e só eu.")
    add_quote("Filipe Delgado", "2022-08-23, 12:42", "Tinhas enviado o vídeo e tinha entrado 5000 para as contas e para uma casa para ti garantida.")
    add_quote("Nuno Duarte", "2022-08-23, 12:19", "tenho as miudas num Hotel.")

    # Secao III
    p_s3 = doc.add_paragraph()
    p_s3.paragraph_format.space_before = Pt(12)
    p_s3.paragraph_format.space_after = Pt(4)
    r_s3 = p_s3.add_run("III. DO DIREITO DE RETENÇÃO E JURISPRUDÊNCIA VINCULATIVA (ART. 754.º DO CC)")
    r_s3.bold = True
    r_s3.font.size = Pt(11)

    p_p3 = doc.add_paragraph()
    p_p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_p3.paragraph_format.first_line_indent = Inches(0.4)
    p_p3.paragraph_format.line_spacing = 1.3
    p_p3.add_run("3. O Réu suportou obras de conservação indispensáveis no montante superior a 120.000,00 €. Como decidiu o ")
    p_p3.add_run("Supremo Tribunal de Justiça (Acórdão de 2011-11-09, Proc. 61/10)").italic = True
    p_p3.add_run(", o direito de retenção é oponível erga omnes e prevalece sobre medidas cautelares ou possessórias (Art. 759.º, n.º 2 do CC), constituindo obstáculo legal à desocupação sem prévio reembolso.")

    # Pedidos
    p_s4 = doc.add_paragraph()
    p_s4.paragraph_format.space_before = Pt(12)
    p_s4.paragraph_format.space_after = Pt(4)
    r_s4 = p_s4.add_run("IV. DOS PEDIDOS")
    r_s4.bold = True
    r_s4.font.size = Pt(11)

    def add_pedido(letra, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Inches(0.2)
        p.add_run(f"{letra}) ").bold = True
        p.add_run(text)

    add_pedido("a", "Seja julgada procedente a Exceção Dilatória de Preterição de Litisconsórcio Passivo Necessário e o Réu absolvido da instância (Art. 577.º, al. e) do CPC);")
    add_pedido("b", "Caso assim não se entenda, seja a ação julgada improcedente e reconhecido ao Réu o DIREITO DE RETENÇÃO (Artigo 754.º do Código Civil) pelo crédito de 120.000,00 €;")
    add_pedido("c", "Seja a Autora condenada em custas e como litigante de má-fé em multa e indemnização condigna.")

    # Fecho
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sn = p_sign.add_run("O Réu / Mandatário:\n\n__________________________________________________\n(Nuno Miguel Silva Duarte)")
    r_sn.bold = True

    doc.save(str(DOCX_OUT))
    print(f"[+] DOCX gerado: {DOCX_OUT}")

def gerar_modelo_pdf():
    styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        "CitiusHeader", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=9.5, leading=12.5, textColor=colors.HexColor("#1e293b")
    )
    title_style = ParagraphStyle(
        "CitiusTitle", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#0f172a"), alignment=1, spaceAfter=8
    )
    body_style = ParagraphStyle(
        "CitiusBody", parent=styles["Normal"], fontName="Times-Roman", fontSize=9.5, leading=13.5, textColor=colors.HexColor("#000000"), alignment=4
    )
    quote_style = ParagraphStyle(
        "QuoteStyle", parent=body_style, fontName="Times-Italic", fontSize=8.5, leading=11.5, textColor=colors.HexColor("#334155")
    )

    doc = SimpleDocTemplate(str(PDF_OUT), pagesize=A4, rightMargin=1.8*cm, leftMargin=1.8*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    story = []

    # Cabecalho
    hdr_data = [
        [
            Paragraph("<b>TRIBUNAL JUDICIAL DA COMARCA DE LISBOA</b><br/>Juízo Central Cível de Lisboa<br/><font size=7.5 color='#64748b'>Palácio da Justiça — Lisboa</font>", header_style),
            Paragraph("<b>PROCESSO:</b> 15547/26.0T8LSB<br/><b>ESPÉCIE:</b> Ação de Reivindicação<br/><b>VALOR:</b> 125.000,00 €", header_style)
        ]
    ]
    t_hdr = Table(hdr_data, colWidths=[11*cm, 6.5*cm])
    t_hdr.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LINEBELOW', (0,0), (-1,-1), 1, colors.HexColor("#94a3b8")),
    ]))
    story.append(t_hdr)
    story.append(Spacer(1, 8))

    story.append(Paragraph("<b>CONTESTAÇÃO E RECONVENÇÃO</b>", title_style))
    story.append(Paragraph("<b>(Preterição de Litisconsórcio, Falsidade de Faturas e Direito de Retenção)</b>", header_style))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=8))

    story.append(Paragraph("<b>NUNO MIGUEL SILVA DUARTE</b>, Réu nos autos à margem identificados em que é Autora <b>MARIA TERESA CASTRO BANGUESES RIBEIRO</b>, vem apresentar a sua Contestação:", body_style))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>I. EXCEÇÃO DILATÓRIA: PRETERIÇÃO DE LITISCONSÓRCIO NECESSÁRIO (ART. 33.º CPC)</b>", header_style))
    story.append(Paragraph("1. A Autora celebrou o Contrato de Arrendamento N.º 1195528 com a sociedade <b>LISBON EXPERIENCE - ADMINISTRAÇÃO DE IMÓVEIS, LDA.</b> (conforme prova <code>LISTA_CONTRATOS_TERESA.xls</code>). A falta de chamada da sociedade locatária determina a ilegitimidade passiva singular e a absolvição da instância (Art. 577.º, al. e) do CPC).", body_style))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>II. CONFISSÕES ESCRITAS E PROVA DOCUMENTAL (WHATSAPP E FATURA 1000002)</b>", header_style))
    story.append(Paragraph("2. As mensagens oficiais trocadas com a gerência provam a exclusividade de responsabilidade e a retenção de fundos devidos ao Réu:", body_style))
    story.append(Spacer(1, 4))

    quote_data = [
        [Paragraph("<b>[2023-01-24, 17:15] Filipe Delgado:</b> <i>'O único responsável sou eu. Documento foi assinado por mim como confissão de dívida e autenticado pelo advogado. Eu e só eu.'</i>", quote_style)],
        [Paragraph("<b>[2022-08-23, 12:42] Filipe Delgado:</b> <i>'Tinhas enviado o vídeo e tinha entrado 5000 para as contas e para uma casa para ti garantida.'</i>", quote_style)],
        [Paragraph("<b>[2022-08-23, 12:19] Nuno Duarte:</b> <i>'tenho as miudas num Hotel.'</i>", quote_style)]
    ]
    t_q = Table(quote_data, colWidths=[17.5*cm])
    t_q.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 6)
    ]))
    story.append(t_q)
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>III. DIREITO DE RETENÇÃO E JURISPRUDÊNCIA STJ (ART. 754.º DO CC)</b>", header_style))
    story.append(Paragraph("3. O Réu realizou benfeitorias necessárias superiores a 120.000,00 €. Como decidiu o <b>Supremo Tribunal de Justiça (Acórdão de 2011-11-09, Proc. 61/10)</b>, o direito de retenção é oponível <i>erga omnes</i> e prevalece sobre medidas cautelares (Art. 759.º, n.º 2 do CC).", body_style))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>IV. PEDIDOS:</b> Absolvição da instância por ilegitimidade; improcedência da ação; e procedência da Reconvenção com reconhecimento do <b>Direito de Retenção pelo crédito de 120.000,00 €</b>.", header_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("O Réu: __________________________________________________<br/>(NUNO MIGUEL SILVA DUARTE)", header_style))

    doc.build(story)
    print(f"[+] PDF gerado: {PDF_OUT}")

if __name__ == "__main__":
    gerar_modelo_docx()
    gerar_modelo_pdf()
