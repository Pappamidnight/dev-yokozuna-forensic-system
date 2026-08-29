#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gerar_template_docx_libreoffice.py - Gerador de Peças e Minutas Formatadas em DOCX para LibreOffice Writer / Word.
Gera documentos editáveis com tipografia jurídica, cabeçalho do Tribunal, citações destacadas e assinatura.
"""

import os
import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DEV_ROOT = Path(r"C:\Users\Yokozuna\Dev")
OUTPUT_DIR = DEV_ROOT / "OUTPUT_CENTRALIZADO" / "05_PDFS_GERADOS_PARA_IMPRESSAO"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUTPUT_DIR / "MINUTA_CONTESTACAO_15547_LIBREOFFICE.docx"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def gerar_minuta_docx():
    print("=" * 80)
    print(" A GERAR MINUTA JUDICIAL EDITÁVEL EM DOCX PARA LIBREOFFICE WRITER")
    print(f" Destino: {DOCX_PATH}")
    print("=" * 80)

    doc = docx.Document()

    # Configuração de Margens (Padrão Forense Tribunal)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(0.8)

    # Estilo Normal
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    # 1. Cabeçalho Oficial do Tribunal
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    cell_l, cell_r = header_table.rows[0].cells
    cell_l.width = Inches(4.0)
    cell_r.width = Inches(2.5)

    p_l = cell_l.paragraphs[0]
    r1 = p_l.add_run("TRIBUNAL JUDICIAL DA COMARCA DE LISBOA\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p_l.add_run("Juízo Central Cível de Lisboa\nPalácio da Justiça — Lisboa")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = p_r.add_run("PROCESSO N.º 15547/26.0T8LSB\n")
    r3.bold = True
    r3.font.size = Pt(10)
    r4 = p_r.add_run("Espécie: Ação de Reivindicação\nValor: 125.000,00 €")
    r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 2. Endereçamento ao Juiz
    p_juiz = doc.add_paragraph()
    r_j = p_juiz.add_run("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DO JUÍZO CENTRAL CÍVEL DE LISBOA")
    r_j.bold = True
    r_j.font.size = Pt(11)
    p_juiz.paragraph_format.space_after = Pt(14)

    # 3. Título da Peça
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("CONTESTAÇÃO E RECONVENÇÃO\n(COM ARGUIÇÃO DE DIREITO DE RETENÇÃO E PRETERIÇÃO DE LITISCONSÓRCIO)")
    r_t.bold = True
    r_t.font.size = Pt(13)
    p_title.paragraph_format.space_after = Pt(18)

    # 4. Qualificação do Réu
    p_qual = doc.add_paragraph()
    p_qual.paragraph_format.line_spacing = 1.3
    p_qual.paragraph_format.space_after = Pt(12)
    p_qual.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_qual.add_run("NUNO MIGUEL SILVA DUARTE").bold = True
    p_qual.add_run(", NIF 254048382, Réu nos autos da Ação de Processo Comum supra identificada em que é Autora ")
    p_qual.add_run("MARIA TERESA CASTRO BANGUESES RIBEIRO").bold = True
    p_qual.add_run(", vem mui respeitosamente apresentar a sua CONTESTAÇÃO, nos termos e com os fundamentos seguintes:")

    # 5. Secções de Defesa
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    def add_body_paragraph(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold_prefix:
            p.add_run(bold_prefix).bold = True
        p.add_run(text)

    def add_quote_box(speaker, date_time, quote_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.rows[0].cells[0]
        cell.width = Inches(6.0)
        set_cell_background(cell, "F1F5F9")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        
        r_hdr = p.add_run(f"[{date_time}] {speaker}:\n")
        r_hdr.bold = True
        r_hdr.font.size = Pt(9.5)
        r_hdr.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        
        r_txt = p.add_run(f"\"{quote_text}\"")
        r_txt.italic = True
        r_txt.font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # I. EXCEÇÃO DILATÓRIA
    add_section_header("I. POR EXCEÇÃO DILATÓRIA: DA PRETERIÇÃO DE LITISCONSÓRCIO NECESSÁRIO (ART. 33.º CPC)")
    add_body_paragraph("A Autora intentou a presente ação apenas contra o Réu singular, omitindo deliberadamente a sociedade locatária LISBON EXPERIENCE - ADMINISTRAÇÃO DE IMÓVEIS, LDA., com a qual celebrou o Contrato de Arrendamento N.º 1195528 (Matriz 110661-U-231-4), pelo que se verifica a ilegitimidade passiva singular e a falta de litisconsórcio passivo necessário.", "1. ")

    # II. IMPUGNAÇÃO DAS FATURAS
    add_section_header("II. DA FALSIDADE MATERIAL DAS FATURAS DO PRÉDIO CONTÍGUO (PRÉDIO 31 vs 33)")
    add_body_paragraph("As faturas e balancetes juntos pela Autora na petição inicial respeitam a despesas do prédio vizinho (Rua da Palmeira n.º 31 / Matriz 110661-U-229) e não à fração do 4.º andar da Rua da Palmeira n.º 33, configurando indução dolosa do Tribunal em erro e litigância de má-fé.", "2. ")

    # III. DIREITO DE RETENÇÃO E BENFEITORIAS
    add_section_header("III. DO DIREITO DE RETENÇÃO PELO CRÉDITO DE BENFEITORIAS (ART. 754.º DO CÓDIGO CIVIL)")
    add_body_paragraph("O Réu suportou obras de conservação indispensáveis no montante superior a 120.000,00 €, detendo direito legal de retenção sobre a fração até ao integral reembolso das benfeitorias.", "3. ")

    # IV. CORTE DE ÁGUA E COAÇÃO
    add_section_header("IV. DA PROVA DOCUMENTAL DO CORTE DE ÁGUA E COAÇÃO HABITACIONAL")
    add_body_paragraph("A confissão expressa em comunicações escritas de 23/08/2022 demonstra que a fração foi privada do abastecimento público de água durante meses:", "4. ")
    
    add_quote_box("Filipe Delgado", "23/08/2022, 12:42", "Tinhas enviado o vídeo e tinha entrado 5000 para as contas e para uma casa para ti garantida")
    add_quote_box("Nuno Duarte", "23/08/2022, 12:19", "tenho as miudas num Hotel")
    add_quote_box("Filipe Delgado", "23/08/2022, 19:21", "Garantia da tua casa para viver só e possível com renda da sky e equilibrar as contas")

    # V. PEDIDOS
    add_section_header("V. DOS PEDIDOS")
    add_body_paragraph("Seja julgada procedente a exceção dilatória de preterição de litisconsórcio necessário e o Réu absolvido da instância (Art. 577.º, al. e) do CPC);", "a) ")
    add_body_paragraph("Caso assim não se entenda, seja a ação julgada totalmente improcedente e reconhecido ao Réu o Direito de Retenção (Art. 754.º CC) pelo crédito de benfeitorias;", "b) ")
    add_body_paragraph("Seja a Autora condenada em custas e como litigante de má-fé.", "c) ")

    # Assinatura
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_s = p_sign.add_run("O Réu / Mandatário:\n\n__________________________________________________\n(Nuno Miguel Silva Duarte)")
    r_s.bold = True

    doc.save(str(DOCX_PATH))
    print(f"[+] Minuta em DOCX gravada com sucesso em: {DOCX_PATH}")
    print("=" * 80)

if __name__ == "__main__":
    gerar_minuta_docx()
